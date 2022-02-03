# pip install python-docx
# pip install pillow

from docx import Document
from docx.shared import Inches
from PIL import Image, UnidentifiedImageError
from os import listdir, path
from sys import argv
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

directory = '.'
if len(argv) > 1:
    directory = argv[1]

document = Document()
image_max_width = Inches(6.1)
image_max_height = Inches(7.6)
# image_max_height = Inches(8.6)
breakpoint_aspect_ratio = image_max_width / image_max_height

files = listdir(directory)

image_paths = []
rel_image_paths = []
pil_images = []
image_bytes = []

for file in files:
    abs_file = path.join(directory, file)
    try:
        pil_images.append(Image.open(abs_file))
        image_paths.append(abs_file)
        rel_image_paths.append(file)
        f = io.BytesIO()
        Image.open(abs_file).convert('RGB').save(f, format='JPEG')
        image_bytes.append(f)
    except Exception as e:
        print("Exception: " + str(e))

for index in range(len(image_paths)):
    section = document.sections[index]
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].text = "\t\tGlobale ID"
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = "Lokale ID\t" + rel_image_paths[index] + "\tSeite " + str(index + 1)
    image_path = image_paths[index]
    img = pil_images[index]
    image_aspect_ratio = img.width / img.height
    if image_aspect_ratio > breakpoint_aspect_ratio:
        document.add_picture(image_bytes[index], width=image_max_width)
    else:
        image = document.add_picture(image_bytes[index], height=image_max_height)
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = document.add_paragraph("Bild " + str(index + 1))
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    for i in range(4):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        
    if index != len(image_paths) - 1:
        document.add_page_break()
        document.add_section()

document.save('output.docx')
